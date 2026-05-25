import os
import docx

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        # Try replacing inside runs to preserve formatting
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return
        # Fallback if text spans runs
        paragraph.text = paragraph.text.replace(old_text, new_text)

def replace_text_in_cell(cell, old_text, new_text):
    for p in cell.paragraphs:
        replace_text_in_paragraph(p, old_text, new_text)

def apply_corrections(docx_path):
    print(f"Applying corrections to {docx_path}...")
    doc = docx.Document(docx_path)
    
    # 1. Apply paragraph replacements
    for p in doc.paragraphs:
        # Correction 1: Link replacement
        replace_text_in_paragraph(p, 
            "https://github.com/igeoo/scipra-ai-mcdm-mineral-policy-africa.git", 
            "https://github.com/igeoo/aiml-mrms-policy-prediction.git"
        )
        # General replacements
        replace_text_in_paragraph(p, "LP/MILP", "Continuous fractional LP")
        replace_text_in_paragraph(p, "feature importance signals. (t)", "feature importance signals f_i(t)")
        replace_text_in_paragraph(p, "(t) from AI layer", "f_i(t) from AI layer")
        replace_text_in_paragraph(p, "scores; budget B", "TOPSIS closeness scores; budget B")

    # 2. Apply table replacements
    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                # Correction 1: Link in tables (if any)
                replace_text_in_cell(cell, 
                    "https://github.com/igeoo/scipra-ai-mcdm-mineral-policy-africa.git", 
                    "https://github.com/igeoo/aiml-mrms-policy-prediction.git"
                )
                # Correction 2: LP/MILP -> Continuous fractional LP
                replace_text_in_cell(cell, "LP/MILP", "Continuous fractional LP")
                # Correction 4-6: Table 2 placeholders
                replace_text_in_cell(cell, "feature importance signals. (t)", "feature importance signals f_i(t)")
                replace_text_in_cell(cell, "(t) from AI layer", "f_i(t) from AI layer")
                replace_text_in_cell(cell, "scores; budget B", "TOPSIS closeness scores; budget B")
                # Correction 7-8: Table 5 placeholders
                replace_text_in_cell(cell, "Aggregated: / / /", "Aggregated: EV / RC / EI / OF")
                replace_text_in_cell(cell, "Aggregated: / /  /", "Aggregated: EV / RC / EI / OF")
                replace_text_in_cell(cell, "0.254 / 0.217 / 0.289 / 0.240", "0.148 / 0.351 / 0.174 / 0.326")

    # 3. Correction 3: Add Table 8 calibrated-assumption footnote
    # Table 8 is index 8 (9th table in document)
    if len(doc.tables) > 8:
        table_8 = doc.tables[8]
        footnote_text = "Note: Calibration assumptions are based on WGI 2022 and Fraser PPI 2022 normalized baseline scores."
        
        # Insert a paragraph after Table 8
        paragraph_element = table_8._element.addnext(docx.oxml.OxmlElement('w:p'))
        p = docx.text.paragraph.Paragraph(paragraph_element, table_8._parent)
        run = p.add_run(footnote_text)
        run.font.size = docx.shared.Pt(9)
        run.italic = True
        print("Added Table 8 footnote.")
    else:
        print("Warning: Table 8 not found in document.")

    doc.save(docx_path)
    print("Corrections applied successfully.")

if __name__ == "__main__":
    docx_file = "../../AIML_MRMS_07052026.docx"
    if not os.path.exists(docx_file):
        # Try local folder
        docx_file = "AIML_MRMS_07052026.docx"
    apply_corrections(docx_file)
