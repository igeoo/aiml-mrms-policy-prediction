"""Generate the Version C candidate tables and consolidated supervisor report.

This script does not rerun the statistical analyses. It consumes the locked,
validated artifacts and creates publication-facing candidate tables plus a
Word brief that records claims, limitations, and decisions still requiring
author approval.
"""

from __future__ import annotations

import csv
import hashlib
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ROBUST = ROOT / "audit_artifacts" / "20260717_six_dimension_robustness_final"
FRASER = ROOT / "audit_artifacts" / "20260717_fraser_convergent_validation"
MODELING = ROOT / "audit_artifacts" / "20260717_novelty_modeling"
TABLE_DIR = ROOT / "results" / "tables"
OUTPUT_DIR = ROOT / "output" / "doc"
EVIDENCE_DIR = ROOT / "audit_artifacts" / "20260717_version_c_acceptance_package"


REFERENCES = [
    {
        "area": "Composite indicators",
        "source": "OECD, European Union and EC-JRC (2008), Handbook on Constructing Composite Indicators",
        "url": "https://doi.org/10.1787/9789264043466-en",
        "prior_art": "Weighting, aggregation, uncertainty and sensitivity analysis for country rankings are established practice.",
        "boundary": "Monte Carlo analysis alone is not the paper's novelty.",
    },
    {
        "area": "Governance measurement",
        "source": "Kaufmann and Kraay (2024), The Worldwide Governance Indicators: Methodology and 2024 Update",
        "url": "https://doi.org/10.1596/1813-9450-10952",
        "prior_art": "The WGI already estimate six governance dimensions and report measurement uncertainty.",
        "boundary": "The proposed PCI/RPCI aggregate is an author-defined analytical layer, not an official World Bank index.",
    },
    {
        "area": "Revised WGI data",
        "source": "World Bank (2025), Worldwide Governance Indicators documentation and revised methodology",
        "url": "https://www.worldbank.org/en/publication/worldwide-governance-indicators/documentation",
        "prior_art": "The official series supplies six 0-100 dimension scores, standard errors and confidence intervals.",
        "boundary": "Short-run changes and overlapping intervals must not be over-interpreted.",
    },
    {
        "area": "Mining policy benchmark",
        "source": "Fraser Institute (2023/2024), Annual Survey of Mining Companies",
        "url": "https://www.fraserinstitute.org/studies/annual-survey-of-mining-companies-2023",
        "prior_art": "PPI is an established perception-based mining-policy benchmark.",
        "boundary": "PPI is used only for convergent validation, never as both a feature and target.",
    },
    {
        "area": "Mining ranking limitations",
        "source": "The perils of ranking mining countries and regions (Mineral Economics, 2024)",
        "url": "https://doi.org/10.1007/s13563-023-00405-y",
        "prior_art": "Mining jurisdiction rankings can obscure survey uncertainty and contested policy concepts.",
        "boundary": "Results are screening evidence, not precise prescriptions or causal effects.",
    },
    {
        "area": "Governance in mining risk",
        "source": "Kuehnel et al. (2023), Correlation analysis of country governance indicators and mining incidents",
        "url": "https://doi.org/10.1016/j.resourpol.2023.103762",
        "prior_art": "Country governance indicators have been tested in mining contexts and can fail for specific outcomes.",
        "boundary": "Construct validity is outcome-specific; the paper reports its failed general-FDI test.",
    },
    {
        "area": "Mining MCDM",
        "source": "Support of mining investment choice decisions with the use of multi-criteria method",
        "url": "https://www.sciencedirect.com/science/article/abs/pii/S0301420716303348",
        "prior_art": "AHP-based mining investment choice is established.",
        "boundary": "AHP or TOPSIS alone cannot be claimed as novel.",
    },
    {
        "area": "Strategic-mineral indicators",
        "source": "Synthesized indicator for evaluating security of strategic minerals: lithium case study (2020)",
        "url": "https://www.sciencedirect.com/science/article/pii/S0301420720309466",
        "prior_art": "WGI and Fraser PPI have already appeared in composite mineral-security indicators.",
        "boundary": "The contribution must be the validated integration and uncertainty discipline, not merely data-source combination.",
    },
]


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, value: object, bold: bool = False, color: str | None = None, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None, font_size: float = 8.5):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF", size=font_size)
        shade(table.rows[0].cells[index], "1F4E78")
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_values in rows:
        cells = table.add_row().cells
        row_properties = table.rows[-1]._tr.get_or_add_trPr()
        prevent_split = OxmlElement("w:cantSplit")
        row_properties.append(prevent_split)
        for index, value in enumerate(row_values):
            set_cell_text(cells[index], value, size=font_size)
            if len(table.rows) % 2 == 1:
                shade(cells[index], "EAF2F8")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph()
    return table


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def build_table8() -> pd.DataFrame:
    scores = pd.read_csv(ROBUST / "candidate_table8_six_dimension_2024_matrix.csv")
    uncertainty = pd.read_csv(ROBUST / "wgi_measurement_uncertainty_2024.csv")
    uncertainty = uncertainty[uncertainty["scheme"] == "uniform_primary"]
    frame = scores.merge(uncertainty, on="iso3", validate="one_to_one")
    frame = frame.sort_values("pci_rank__uniform_primary")
    result = pd.DataFrame(
        {
            "rank": frame["pci_rank__uniform_primary"].astype(int),
            "country": frame["country"],
            "iso3": frame["iso3"],
            "pci": frame["pci__uniform_primary"].round(4),
            "pci_measurement_95_ci": frame.apply(lambda row: f"{row.pci_mc_lower95:.4f}-{row.pci_mc_upper95:.4f}", axis=1),
            "pci_rank_measurement_95_interval": frame.apply(lambda row: f"{int(row.pci_rank_lower95)}-{int(row.pci_rank_upper95)}", axis=1),
            "rpci": frame["rpci__uniform_primary"].round(4),
            "rpci_rank": frame["rpci_rank__uniform_primary"].astype(int),
            "rpci_measurement_95_ci": frame.apply(lambda row: f"{row.rpci_mc_lower95:.4f}-{row.rpci_mc_upper95:.4f}", axis=1),
        }
    )
    return result


def build_table9() -> pd.DataFrame:
    stability = pd.read_csv(ROBUST / "candidate_table9_weight_scheme_stability.csv")
    perturbation = load_json(ROBUST / "weight_perturbation_summary.json")
    fraser_metrics = load_json(FRASER / "fraser_logo_metrics.json")
    fraser_boot = load_json(FRASER / "fraser_cluster_bootstrap.json")
    fdi_boot = load_json(MODELING / "cluster_bootstrap_increment.json")
    construct = load_json(MODELING / "construct_validity.json")

    rows: list[dict[str, str]] = []
    names = {
        "pca_sensitivity": "PCA absolute-loading weights vs uniform",
        "entropy_sensitivity": "Entropy weights vs uniform",
        "fraser_bivariate_sensitivity": "Fraser-correlation weights vs uniform",
    }
    for record in stability.to_dict("records"):
        if record["scheme"] == "uniform_primary":
            continue
        rows.append(
            {
                "evidence_block": "Alternative weighting",
                "test": names[record["scheme"]],
                "result": f"PCI rho={record['pci_spearman_vs_uniform']:.3f}; RPCI rho={record['rpci_spearman_vs_uniform']:.3f}; mean |PCI diff|={record['mean_abs_pci_difference']:.4f}",
                "interpretation": "Rank ordering is highly stable; uniform weighting remains the transparent primary specification.",
            }
        )
    rows.extend(
        [
            {
                "evidence_block": "Weight uncertainty",
                "test": "10,000 Dirichlet weight perturbations",
                "result": f"PCI rank rho median={perturbation['pci_rank_spearman_median']:.3f} (2.5%={perturbation['pci_rank_spearman_lower95']:.3f}); RPCI median={perturbation['rpci_rank_spearman_median']:.3f} (2.5%={perturbation['rpci_rank_spearman_lower95']:.3f})",
                "interpretation": "Rank conclusions are robust to broad symmetric weight perturbation.",
            },
            {
                "evidence_block": "Mining convergent validity",
                "test": "WGI governance mean vs Fraser PPI, 2019-2024",
                "result": f"Spearman rho={fraser_boot['estimate']:.3f}; country-cluster bootstrap 95% CI {fraser_boot['cluster_bootstrap_ci95_lower']:.3f}-{fraser_boot['cluster_bootstrap_ci95_upper']:.3f}; n=70, 14 countries",
                "interpretation": "Strong association with a mining-specific policy-perception benchmark; not causal validation.",
            },
            {
                "evidence_block": "Mining out-of-country validation",
                "test": "Nested leave-one-country-out Ridge prediction of Fraser PPI",
                "result": f"R2={fraser_metrics['r2']:.3f}; MAE={fraser_metrics['mae']:.2f}; Spearman rho={fraser_metrics['spearman']:.3f}",
                "interpretation": "Moderate out-of-country generalisation; coefficients are not used as primary weights because dimensions are collinear.",
            },
            {
                "evidence_block": "Falsification / discriminant boundary",
                "test": "Governance plus controls vs controls for next-year general FDI",
                "result": f"Incremental MAE improvement={fdi_boot['estimate']:.4f}; country-cluster bootstrap 95% CI {fdi_boot['ci95_lower']:.4f} to {fdi_boot['ci95_upper']:.4f}; P(improvement)={fdi_boot['bootstrap_probability_improvement']:.3f}",
                "interpretation": "No reliable predictive gain. General FDI is rejected as the success target for this governance layer.",
            },
            {
                "evidence_block": "Construct diagnostics",
                "test": "Six WGI dimensions, pooled 2002-2024",
                "result": f"Cronbach alpha={construct['cronbach_alpha_pooled']:.3f}; PC1={construct['pca_pc1_explained_variance']:.3f}; max VIF={construct['maximum_vif']:.2f}",
                "interpretation": "Strong shared governance signal and high collinearity; this supports parsimonious aggregation but not unidimensionality or causal attribution.",
            },
        ]
    )
    return pd.DataFrame(rows)


def build_evidence_matrix() -> pd.DataFrame:
    proposed = (
        "Construct-separated, uncertainty-aware six-dimensional governance-coherence layer for mineral-investment screening, "
        "with leakage-free mining-specific external validation, explicit null-result reporting, and reproducible rank/weight stress tests."
    )
    return pd.DataFrame([{**record, "defensible_increment": proposed} for record in REFERENCES])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    for style_name, size, color in [("Title", 20, "17365D"), ("Heading 1", 15, "1F4E78"), ("Heading 2", 12, "2F75B5")]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(color)


def add_landscape_section(document: Document):
    section = document.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    return section


def build_document(table8: pd.DataFrame, table9: pd.DataFrame, evidence: pd.DataFrame, output_path: Path) -> None:
    document = Document()
    configure_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AIML-MRMS Version C\nComputational Evidence and Supervisor Decision Brief")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Candidate replacement for the disputed Tables 8 and 9 | 17 July 2026").italic = True

    document.add_heading("Executive conclusion", level=1)
    document.add_paragraph(
        "The former four-configuration performance table should remain withdrawn. The code does not execute four independent decision architectures, "
        "does not contain a closed feedback loop, and does not test convergence. Version C now provides a reproducible alternative: a six-dimensional "
        "country-governance coherence analysis using official WGI absolute scores, explicit measurement uncertainty, alternative-weight sensitivity, "
        "temporal diagnostics, and mining-specific external validation."
    )
    document.add_paragraph(
        "This package is computationally ready for author review. It is not yet a publishable manuscript by itself: the methods, table captions, claims, "
        "limitations and contribution statement must be rewritten to match the analysis. Acceptance cannot be guaranteed, but the known computational "
        "and construct-validity loopholes are now exposed rather than hidden."
    )

    document.add_heading("Direct response to the supervisor's questions", level=1)
    document.add_heading("Question 1 - Is the table comparing the right things?", level=2)
    document.add_paragraph(
        "Answer: no, not as independent architectures. The historical labels were broader than the executed computation. The implementation represents "
        "alternative weight states applied to essentially the same MCDM/PCI calculation. The final state is the result of two predefined convex updates; "
        "there is no convergence check and no country-level TOPSIS, allocation or PCI result fed back into the weights."
    )
    add_table(
        document,
        ["Historical label", "What is actually executed", "Safe description"],
        [
            ["Standalone MCDM", "Base expert/AHP weight vector applied to the common governance inputs", "Base-weight scenario"],
            ["Standalone AI", "Normalised model-derived coefficient signal used as weights; no independent AI decision pipeline", "AI-derived weighting sensitivity"],
            ["Static AI+MCDM", "One convex update from the base vector toward the fixed signal", "First predefined blended-weight state"],
            ["Full AIML-MRMS", "A second convex update toward the same fixed signal", "Second predefined blended-weight state; not converged and not feedback"],
        ],
        [1.35, 4.2, 2.3],
    )
    document.add_paragraph(
        "Decision: do not restore the old comparison merely by substituting new numbers. Version C separates country governance from project-level "
        "EV/RC/EI/OF AHP criteria and replaces the disputed comparison with an uncertainty-aware governance evidence table."
    )

    document.add_heading("Question 2 - Why did the adaptive scenario lose?", level=2)
    document.add_paragraph(
        "The earlier explanation that the CV penalty caused the loss was plausible but initially unproven. The metric is sensitive to both the weighted "
        "mean and dispersion: PCI = weighted mean x (1 - weighted CV), while RPCI additionally applies a weighted-Gini coherence penalty. Concentrated "
        "weights can therefore raise or lower the final score depending on each country's governance profile. This is metric sensitivity, not evidence "
        "that an adaptive architecture succeeds or fails. Because the historical experiment did not implement outcome feedback, the term 'adaptive "
        "performance' should not be used for those results."
    )

    document.add_heading("Why Version C is methodologically different", level=1)
    add_bullets(
        document,
        [
            "Uses all six official WGI dimensions: Voice and Accountability, Political Stability, Government Effectiveness, Regulatory Quality, Rule of Law, and Control of Corruption.",
            "Uses official 0-100 absolute scores divided by 100; it does not re-min-max countries within this 16-country sample.",
            "Keeps country governance weights separate from project-level economic viability, resource capacity, environmental/ESG, and operational feasibility weights.",
            "Uses uniform six-dimension weights as the transparent primary specification, with PCA, entropy and mining-benchmark correlation weights only as sensitivity checks.",
            "Propagates published WGI standard errors through 10,000 simulations and separately perturbs weights through 10,000 Dirichlet draws.",
            "Validates against Fraser PPI in a leakage-free design. Fraser PPI is never used simultaneously as both a model input and its target.",
            "Reports a failed general-FDI predictive test as a boundary condition: FDI is not treated as proof of governance-layer success.",
        ],
    )

    document.add_heading("Computational validation gates already passed", level=2)
    add_bullets(
        document,
        [
            "Complete official WGI panel: 368 country-year observations, covering 16 study countries and 23 years (2002-2024), with six scores and uncertainty fields.",
            "No within-sample min-max rescaling and no hard-coded country performance constants in the candidate tables.",
            "Leakage removed: the mining-policy benchmark is an external target/validation measure, not a simultaneous predictor.",
            "Country-held-out and temporal validation designs are explicit; bootstrap resampling is clustered by country.",
            "All primary and sensitivity weights sum to one; 2024 score and uncertainty outputs are complete for all 16 countries.",
            "Independent repeat of the robustness run produced 11 byte-identical files.",
        ],
    )
    document.add_heading("What remains outside Version C", level=2)
    document.add_paragraph(
        "Version C does not manufacture a feedback loop, estimate project cash flows, or supply missing geology and firm-level investment decisions. "
        "Those would be separate extensions requiring new data and validation. Its role is narrower and defensible: country-governance screening and "
        "coherence diagnostics that can sit beside, but not substitute for, the project-level AHP-TOPSIS analysis."
    )

    add_landscape_section(document)
    document.add_heading("Candidate Table 8", level=1)
    document.add_paragraph(
        "Suggested title: 2024 Six-Dimension Governance Coherence Scores under the Uniform Primary Specification. "
        "Intervals propagate official WGI dimension standard errors under an independent-normal approximation; rank intervals are not significance groups."
    )
    table8_rows = []
    for row in table8.itertuples(index=False):
        table8_rows.append([row.rank, row.country, row.iso3, f"{row.pci:.4f}", row.pci_measurement_95_ci, row.pci_rank_measurement_95_interval, f"{row.rpci:.4f}", row.rpci_rank, row.rpci_measurement_95_ci])
    add_table(
        document,
        ["PCI rank", "Country", "ISO3", "PCI", "PCI 95% interval", "PCI rank interval", "RPCI", "RPCI rank", "RPCI 95% interval"],
        table8_rows,
        [0.58, 1.45, 0.55, 0.62, 1.25, 1.05, 0.62, 0.72, 1.25],
        7.6,
    )
    document.add_paragraph(
        "Interpretation: the point ranking is descriptive. Botswana is first and the Democratic Republic of Congo is last in the 2024 study sample, "
        "but several middle-country measurement-based rank intervals overlap. Policy conclusions should therefore emphasize profiles and broad tiers, not small rank gaps."
    )

    document.add_heading("Candidate Table 9", level=1)
    document.add_paragraph("Suggested title: Robustness, Construct Diagnostics and External Validation of the Six-Dimension Governance Coherence Layer.")
    add_table(
        document,
        ["Evidence block", "Test", "Result", "Interpretation"],
        [[row.evidence_block, row.test, row.result, row.interpretation] for row in table9.itertuples(index=False)],
        [1.25, 2.0, 3.0, 3.7],
        7.5,
    )

    document.add_heading("Novelty claim: what can and cannot be said", level=1)
    document.add_paragraph(
        "Defensible proposed contribution: an uncertainty-aware, construct-separated six-dimensional governance-coherence layer for mineral-investment "
        "screening, supported by leakage-free mining-specific validation, explicit falsification against general FDI, and reproducible temporal, rank "
        "and weighting stress tests. The novelty is the disciplined integration and validation protocol, not any single component."
    )
    add_table(
        document,
        ["Prior-art area", "What already exists", "Claim boundary"],
        [[row.area, row.prior_art, row.boundary] for row in evidence.itertuples(index=False)],
        [1.4, 4.0, 4.4],
        7.8,
    )
    document.add_paragraph("Claims that must not appear:")
    add_bullets(
        document,
        [
            "AHP, TOPSIS, WGI, Fraser PPI, PCA, entropy weighting, machine learning, or Monte Carlo analysis is novel by itself.",
            "The empirical pipeline implements a closed adaptive feedback loop or has converged.",
            "The governance layer causes investment, predicts mining investment generally, or guarantees superior PCI/RPCI.",
            "Small differences in country rank are statistically meaningful when WGI measurement intervals overlap.",
            "Fraser PPI is ground truth. It is a perception benchmark with sampling and conceptual limitations.",
        ],
    )

    document.add_heading("Known limitations retained transparently", level=1)
    add_bullets(
        document,
        [
            "Only 16 study countries are ranked; the results are not an Africa-wide census.",
            "Fraser validation covers 70 country-year observations across 14 countries from 2019-2024; Nigeria and Sierra Leone lack observations and are not imputed.",
            "WGI errors are sampled independently because cross-dimension error covariance is unavailable. This approximation is stated, not concealed.",
            "WGI dimensions are highly collinear. Multivariate coefficient signs are unstable, so predictive coefficients are not promoted as causal or primary weights.",
            "PCI and RPCI are author-defined summaries. Their substantive interpretation and penalty choices require equations, rationale and sensitivity reporting in the manuscript.",
            "Country governance screening cannot replace project-level geology, finance, ESG due diligence, community consent or legal review.",
        ],
    )

    document.add_heading("Decisions requested from the supervisor/author", level=1)
    add_bullets(
        document,
        [
            "Approve replacing the historical four-configuration Tables 8 and 9 with the candidate governance-coherence and validation tables in this brief.",
            "Approve uniform weighting across all six WGI dimensions as primary, with three alternative schemes reported only as sensitivity analyses.",
            "Approve the construct boundary: project-level EV/RC/EI/OF remains in the AHP-TOPSIS layer and is not projected onto country-governance dimensions.",
            "Approve terminology: 'governance-coherence layer' or 'evidence-informed governance layer'; remove empirical claims of feedback, convergence and superiority.",
            "Approve the proposed novelty statement subject to a targeted journal-specific literature update before resubmission.",
        ],
    )

    document.add_heading("Manuscript work after approval", level=1)
    add_table(
        document,
        ["Manuscript location", "Required change"],
        [
            ["Abstract and contributions", "Replace superiority/adaptive-performance claims with the validated integration, uncertainty and construct-separation contribution."],
            ["Conceptual framework", "Draw a hard boundary between country governance screening and project-level AHP-TOPSIS criteria."],
            ["Data", "Document official 2025 revised WGI series, all six dimensions, 2002-2024 panel, Fraser coverage and missingness."],
            ["Methods", "Add exact PCI/RPCI equations, uniform primary weights, alternative schemes, WGI-error simulation, weight perturbation, LOCO and bootstrap protocols."],
            ["Results", "Insert approved Tables 8 and 9; report both Fraser convergence and the null general-FDI result."],
            ["Discussion", "Discuss rank uncertainty, collinearity, perception-data limits, non-causality and screening-only use."],
            ["Title/terminology", "Avoid claiming an empirically executed adaptive feedback system unless a real outcome-fed update loop is later implemented and tested."],
            ["Repository statement", "Publish code, input provenance, seeds, manifests, checksums and machine-readable table outputs."],
        ],
        [2.0, 7.8],
        8.2,
    )

    document.add_heading("Reproducibility record", level=1)
    document.add_paragraph(
        "The official WGI panel contains 368 complete country-year rows (16 countries x 23 years, 2002-2024). The robustness engine was run twice; "
        "all 11 generated files were byte-identical. Random seeds are fixed in the analysis scripts. The source workbook SHA256 recorded during extraction is "
        "25a2f9eabb90b0092973392c0b31571aa58b691cc5786292e504b52f693e1eb8."
    )
    add_table(
        document,
        ["Artifact", "SHA256"],
        [
            ["Candidate Table 8 CSV", sha256(TABLE_DIR / "table8_version_c_governance_coherence_2024.csv")],
            ["Candidate Table 9 CSV", sha256(TABLE_DIR / "table9_version_c_validation_robustness.csv")],
            ["Six-dimension robustness manifest", sha256(ROBUST / "six_dimension_robustness_manifest.json")],
            ["Fraser validation manifest", sha256(FRASER / "fraser_validation_manifest.json")],
            ["General-FDI modeling manifest", sha256(MODELING / "modeling_manifest.json")],
        ],
        [3.3, 6.5],
        7.5,
    )

    document.add_heading("References and evidence links", level=1)
    for index, record in enumerate(REFERENCES, start=1):
        document.add_paragraph(f"{index}. {record['source']}. {record['url']}")

    document.add_paragraph()
    final_note = document.add_paragraph()
    final_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final_note.add_run("Prepared as a computational decision brief. Final table numbering and manuscript adoption remain at the author/supervisor's discretion.")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4E78")
    document.save(output_path)


def write_markdown(table8: pd.DataFrame, table9: pd.DataFrame, evidence: pd.DataFrame, path: Path) -> None:
    def markdown_table(frame: pd.DataFrame) -> str:
        values = frame.fillna("").astype(str)
        headers = [str(column).replace("|", "\\|") for column in values.columns]
        output = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
        for record in values.itertuples(index=False, name=None):
            output.append("| " + " | ".join(value.replace("|", "\\|").replace("\n", " ") for value in record) + " |")
        return "\n".join(output)

    lines = [
        "# Version C acceptance package",
        "",
        "## Locked conclusion",
        "",
        "The historical four-configuration table is not a valid comparison of independently executed architectures. The candidate replacement is an uncertainty-aware six-dimensional governance-coherence analysis.",
        "",
        "## Defensible novelty",
        "",
        "Construct-separated, uncertainty-aware governance screening for mineral-investment decision support, with mining-specific leakage-free validation, explicit falsification against general FDI, and reproducible temporal/rank/weight stress tests.",
        "",
        "## Candidate Table 8",
        "",
        markdown_table(table8),
        "",
        "## Candidate Table 9",
        "",
        markdown_table(table9),
        "",
        "## Novelty evidence matrix",
        "",
        markdown_table(evidence[["area", "source", "prior_art", "boundary", "url"]]),
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)

    table8 = build_table8()
    table9 = build_table9()
    evidence = build_evidence_matrix()

    table8_path = TABLE_DIR / "table8_version_c_governance_coherence_2024.csv"
    table9_path = TABLE_DIR / "table9_version_c_validation_robustness.csv"
    evidence_path = EVIDENCE_DIR / "novelty_evidence_matrix.csv"
    report_path = OUTPUT_DIR / "AIML_MRMS_Version_C_Comprehensive_Supervisor_Report.docx"
    markdown_path = EVIDENCE_DIR / "version_c_acceptance_package.md"

    table8.to_csv(table8_path, index=False, quoting=csv.QUOTE_MINIMAL)
    table9.to_csv(table9_path, index=False, quoting=csv.QUOTE_MINIMAL)
    evidence.to_csv(evidence_path, index=False, quoting=csv.QUOTE_MINIMAL)
    build_document(table8, table9, evidence, report_path)
    write_markdown(table8, table9, evidence, markdown_path)

    manifest = {
        "status": "PASS",
        "purpose": "Candidate replacement Tables 8 and 9 and consolidated supervisor brief",
        "author_approval_required": True,
        "outputs": {
            str(path.relative_to(ROOT)): sha256(path)
            for path in [table8_path, table9_path, evidence_path, markdown_path, report_path]
        },
        "claims": {
            "adaptive_feedback_executed": False,
            "causal_investment_effect": False,
            "general_fdi_predictive_validation": False,
            "mining_policy_convergent_validation": True,
            "wgi_measurement_uncertainty_propagated": True,
        },
    }
    (EVIDENCE_DIR / "acceptance_package_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
