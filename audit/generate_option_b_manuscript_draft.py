"""Create a non-destructive Option B revision draft of the active manuscript.

The original manuscript is never edited. This script produces a separately
named draft under the workspace output/doc directory and uses the latest
Option B audit package for Tables 8 and 9.
"""

from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parents[1]
SOURCE = ROOT / "AIML-MRMS_IJSE_03072024_recent_consistency_cleaned.docx"
OUTPUT = WORKSPACE / "output" / "doc" / "AIML_MRMS_Option_B_Revision_Draft_v2.docx"


def latest_option_b_dir() -> Path:
    candidates = sorted((ROOT / "audit_artifacts").glob("*_option_b"))
    if not candidates:
        raise FileNotFoundError("No Option B audit package found")
    return candidates[-1]


def set_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def set_cell(cell, text: str, bold: bool = False, size: int = 7) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def remove_table(table) -> None:
    table._element.getparent().remove(table._element)


def insert_table_after(paragraph, table) -> None:
    paragraph._p.addnext(table._tbl)


def fmt_pair(pci: float, rpci: float) -> str:
    return f"{pci:.3f} / {rpci:.3f}"


def main() -> None:
    option_b = latest_option_b_dir()
    table8 = pd.read_csv(option_b / "table8_manuscript_matrix.csv").sort_values("iso3")
    table9 = pd.read_csv(option_b / "table9_manuscript_sensitivity_summary.csv").sort_values(["alpha", "cycle"])
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    table_style = doc.tables[0].style

    # The indices are checked against the active revision draft and only the
    # claims connected to the Option B evidence boundary are replaced.
    replacements = {
        2: "Public-sector investment and regulatory decisions in resource-constrained governance environments are often supported by fragmented analytical approaches in which predictive models, multi-criteria evaluation, and optimization operate independently. This study develops the Adaptive Intelligent Machine Learning-Mineral Resource Management System (AIML-MRMS), a governance-aware decision-support methodology organized as a five-layer architecture. Rather than proposing new analytical algorithms, the study contributes a transparent integration design and an interpretable weighting analysis that combines expert AHP priorities with an aggregate SVM-derived criterion signal. Using publicly available governance and mining-investment data from 16 African countries, the study evaluates how explicitly defined weighting scenarios affect PCI/RPCI values and country rankings. The mineral-resources application is a proof-of-concept demonstration of reproducible weight propagation and sensitivity analysis; it does not claim empirical superiority of a closed outcome-feedback loop.",
        8: "The contribution of this study is methodological rather than algorithmic. AIML-MRMS specifies a governance-aware decision-support architecture linking predictive evidence, multi-criteria evaluation, constrained optimization, and a proposed operational feedback interface. The empirical proof-of-concept evaluates the implemented AHP-SVM weighting component through reproducible score and ranking sensitivity analyses on 16 African mining countries. This distinction is important: the architecture defines how future operational feedback could be organized, whereas the reported computation evaluates fixed weight states derived from a single aggregate SVM signal. The contribution is therefore a transparent, auditable integration design and an evidence-led analysis of weighting sensitivity, with broader domain validation identified as future work.",
        35: "The input layer aggregates six publicly accessible features: five World Bank WGI percentile ranks (Rule of Law, Regulatory Quality, Government Effectiveness, Control of Corruption, Political Stability) and the Fraser Institute Policy Perception Index. All inputs are min-max normalised to [0, 1] to form the state vector S(t). ESG and regulatory signals are treated as first-class features rather than post-hoc filters. Figure 2 illustrates the proposed five-layer architecture and directional data flows. The operational feedback pathway is an architectural design feature; it is not executed in the fixed-weight-state PCI/RPCI sensitivity analysis reported in Section 6.",
        38: "Figure 2. Proposed layered architecture of AIML-MRMS. The feedback path from Output to Input/AI is an intended operational interface for future deployments; the present proof-of-concept evaluates the AHP-SVM weighting states without outcome feedback.",
        46: "The output layer consolidates investment recommendations and system performance indicators (PCI, RPCI). In a future operational deployment, outputs and newly released governance data could inform a state-update process. In the present proof-of-concept, the SVM criterion signal is computed once from the verified feature matrix and the update rule is applied a fixed number of times. Neither TOPSIS outputs, investment allocations, nor observed country outcomes are fed back into the weights used for the reported PCI/RPCI scenarios.",
        47: "Figure 3 maps the proposed inter-layer interfaces. The reported weighting-sensitivity analysis executes the SVM-to-MCDM weight propagation; the Output-to-Input/AI state-transition interface remains a proposed operational extension rather than an empirically evaluated feedback path.",
        49: "Figure 3. Proposed formal inter-layer interfaces. SVM feature-importance signals are passed to the MCDM weight update. TOPSIS scores can inform investment allocation under the stated optimization model. The Output-to-Input/AI state transition is a future operational interface and is not used to generate the Table 8/9 weighting-sensitivity results.",
        54: "Input: system state S(t); AHP base weights w0; aggregate SVM-derived criterion signal f. The reported proof-of-concept applies the defined blending rule for a fixed number of weight states. Operational performance feedback is a future extension and is not used to generate Tables 8-9.",
        61: "Step 6 - Operational extension: in a deployment with new outcomes or new governance data, update the system state through the specified state-transition interface. This step is not executed in the reported Table 8/9 analysis.",
        75: "where f(t) is the normalised aggregate SVM feature-importance signal for criterion i and alpha is a stability parameter. When alpha approaches 1, expert judgment dominates; when alpha approaches 0, the aggregate data-derived signal has greater influence. Since w(t) and f(t) are unit-sum vectors, the blending rule preserves a unit-sum weight vector. In this proof-of-concept, f(t) is derived once from the WGI plus Fraser PPI feature coefficients and held constant. The reported weight states are therefore the base AHP state followed by one and two predefined applications of the blending rule. No stopping criterion or convergence test is applied; operational deployment would require new data and a separately specified feedback/update protocol.",
        77: "5.5 Proposed state-update feedback equation",
        79: "The following state-transition equation specifies a proposed operational feedback interface. In a deployment, it would require new governance data, observed decisions, or explicitly measured outcomes. It is not invoked in the fixed-state weighting analysis used to generate Tables 8 and 9.",
        123: "6.5 Weighted PCI/RPCI scenario analysis",
        124: "The revised analysis evaluates sensitivity of the Policy Convergence Index (PCI) and Robust Policy Convergence Index (RPCI) to explicitly defined criterion-weighting scenarios. The original unweighted formulas are retained as the reference case. For each scenario, criterion weights are mapped to the six PCI dimensions, a weighted mean and weighted coefficient-of-variation penalty are computed, and RPCI applies a weighted Gini penalty. Uniform dimension weights reproduce the original unweighted PCI/RPCI values exactly for all 16 countries.",
        126: "For dimension scores s_j and dimension weights v_j that sum to one, the weighted score is mu_w = sum_j v_j s_j; CV_w = sqrt(sum_j v_j(s_j-mu_w)^2)/mu_w; PCI_w = mu_w(1-CV_w). The reference case uses v_j = 1/6. Scenario weights are derived from the base AHP vector, the aggregate SVM criterion signal, one blend of those vectors, or two predefined blends, respectively.",
        127: "The Robust Policy Convergence Index is RPCI_w = PCI_w(1-Gini_w), where Gini_w is calculated using the same dimension-weight vector. PCI/RPCI remain internal governance-alignment indicators and are not external performance-validation metrics.",
        130: "The weighted formulation makes the relationship between the scenario definition and the score explicit. The coefficient-of-variation and Gini penalties are reported separately in the audit material so changes in final scores can be examined rather than attributed to adaptation by assumption.",
        131: "Results are interpreted as a transparent sensitivity analysis: changes may be positive, negative, or neutral depending on country input profiles and the stated weight vector. They do not establish that one architecture universally outperforms another.",
        132: "Table 8. PCI/RPCI under alternative expert- and SVM-informed criterion-weighting scenarios, with the original unweighted reference. Each entry is PCI / RPCI.",
        133: "Table 8 reports all 16 countries under the original-unweighted reference, base-AHP weighting, aggregate-SVM weighting, a one-blend state, and a two-blend state. The two-blend state is a fixed calculation state, not a convergence or outcome-feedback result. Full precision, criterion weights, dimension weights, and mean/CV/Gini decompositions are provided in the reproducibility supplement.",
        134: "The legacy hard-coded gain comparison remains withdrawn. The replacement Table 8 is generated from the documented input vectors and criterion-weight scenarios, and it must be interpreted as score sensitivity rather than comparative performance evidence for independently executed architectures.",
        135: "Future work should evaluate convergent validity against independent external governance instruments and test an operational feedback protocol using new time-indexed data or observed outcomes.",
        136: "6.6 Weighting-sensitivity and ranking-stability analysis",
        137: "Table 9 reports the alpha and fixed-update-count sensitivity analysis. Lower alpha values shift the weight state more rapidly toward the fixed aggregate SVM signal, while higher alpha values retain more of the base-AHP prior. The analysis reports score movement and rank correlation rather than assuming that larger PCI/RPCI values imply a superior method.",
        138: "No empirical architecture-superiority matrix is claimed. The scenarios share the same country inputs and differ in their criterion-weight sources or fixed update counts; no separate AI decision pipeline, TOPSIS feedback, investment feedback, or convergence test is executed for the PCI/RPCI comparison.",
        139: "The evidence therefore supports a bounded conclusion: the AHP-SVM integration produces transparent and reproducible changes in criterion weights, PCI/RPCI, and rankings under stated parameter choices. The significance of those changes depends on the governance decision context and should be assessed alongside the decomposition and robustness results.",
        140: "AIML-MRMS should be interpreted as a proposed architecture with an implemented and tested weighting-sensitivity component, not as validated evidence of numerical dominance over all baseline configurations.",
        117: "The effect of alpha is explored further in Section 6.6.2. Lower alpha values give the fixed aggregate SVM signal greater influence at each update; higher alpha values preserve more of the expert-weight prior. The results are reported as a sensitivity analysis, not a convergence result.",
        151: "Sensitivity analysis was conducted for alpha and fixed update count. Table 9 reports the resulting criterion weights, mean absolute score changes relative to the base-AHP state, and rank correlations relative to that state.",
        152: "Across the sensitivity grid, alpha controls the rate at which the fixed weight state moves toward the aggregate SVM signal. The reported rank correlations and score movements provide a reproducible robustness assessment. They are proof-of-concept sensitivity results, not a convergence test or an externally validated policy-performance gain.",
        154: "The proof-of-concept demonstration uses real, publicly retrievable governance and investment datasets, and reported calculations are reproducible from the cited public sources. Four limitations require explicit acknowledgement. First, the sample size (n=16) limits statistical generalisability. Second, the project-level TOPSIS scores are simulated governance-derived scenarios rather than operational mining-project observations. Third, PCI/RPCI are constructed from governance indicators and are internal indicators rather than independent outcome validation. Fourth, the Table 8/9 analysis uses one aggregate SVM-derived signal and fixed applications of the blending rule; it does not feed TOPSIS results, investment allocations, observed outcomes, or new annual data back into the weights. The findings therefore demonstrate computational reproducibility and weighting sensitivity, not empirical superiority or an evaluated closed feedback loop.",
        166: "Consistent with the Capability Test, the contribution is best stated functionally. AIML-MRMS provides a specified, auditable pathway for connecting governance evidence, structured multi-criteria evaluation, and constrained allocation. The present computation demonstrates the reproducible propagation of an aggregate SVM-derived signal through explicitly defined weighting states. The broader architecture specifies interfaces for future operational feedback, but this paper does not claim to have empirically validated that feedback interface.",
        169: "Organizations may use a one-time SVM-to-MCDM weight coupling when an interpretable evidence-informed weighting update is required. Any claimed benefit from closed outcome feedback, however, requires iterative deployment using new governance data or observed decision outcomes and should be evaluated separately.",
        174: "This study developed and computationally demonstrated AIML-MRMS as a governance-aware decision-support methodology organized through a transparent five-layer architecture. The principal contribution is methodological rather than algorithmic: it specifies how established predictive, evaluative, and optimization methods can be connected through explicit interfaces and traceable evidence flow.",
        175: "The study makes three related contributions. It specifies a governance-aware integration architecture; demonstrates an interpretable AHP-SVM weighting mechanism; and provides a reproducible sensitivity analysis of PCI/RPCI and country-ranking effects under stated parameter choices. The contribution is strengthened by the explicit calculation trail and the reporting of both favourable and unfavourable score movements.",
        176: "The methodology was demonstrated using publicly available governance and mining-investment data from 16 African countries. The computational evaluation establishes internal consistency and reproducibility of the implemented weighting analysis. It does not establish universal score improvement, real-world governance outcomes, independent architecture superiority, convergence, or an empirically evaluated closed outcome-feedback loop.",
        183: "All data used in this study are derived from publicly accessible sources: World Bank Worldwide Governance Indicators 2022, Fraser Institute Annual Survey of Mining Companies 2022, and UNCTAD FDI data. The reproducibility package includes the corrected pipeline, processed data, scenario definitions, weighted PCI/RPCI equations, full Table 8/9 calculations, sensitivity outputs, and audit metadata needed to reproduce the reported values.",
    }
    for index, text in replacements.items():
        set_paragraph(doc.paragraphs[index], text)

    # Replace the previous four-country baseline diagnostics with a compact
    # all-country matrix. Full precision remains in the CSV supplement.
    remove_table(doc.tables[7])
    headers = ["ISO3", "Original", "Base AHP", "SVM signal", "One blend", "Two blends"]
    compact8 = doc.add_table(rows=len(table8) + 1, cols=len(headers))
    compact8.style = table_style
    compact8.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(compact8.rows[0].cells, headers):
        set_cell(cell, header, bold=True, size=7)
    for row_idx, (_, item) in enumerate(table8.iterrows(), start=1):
        values = [
            item["iso3"],
            fmt_pair(item["PCI_original_unweighted_reference"], item["RPCI_original_unweighted_reference"]),
            fmt_pair(item["PCI_standalone_mcdm"], item["RPCI_standalone_mcdm"]),
            fmt_pair(item["PCI_standalone_ai"], item["RPCI_standalone_ai"]),
            fmt_pair(item["PCI_static_ai_mcdm"], item["RPCI_static_ai_mcdm"]),
            fmt_pair(item["PCI_aiml_mrms_full_adaptive"], item["RPCI_aiml_mrms_full_adaptive"]),
        ]
        for cell, value in zip(compact8.rows[row_idx].cells, values):
            set_cell(cell, str(value), size=7)
    insert_table_after(doc.paragraphs[132], compact8)

    # Table 9 is a compact robustness summary; country-level values stay in
    # the reproducibility supplement so the manuscript remains readable.
    heading = doc.paragraphs[152]
    table9_title = heading.insert_paragraph_before("Table 9. Sensitivity of weight states, scores, and country rankings to alpha and fixed update count.")
    table9_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headers9 = ["alpha", "state", "EV", "RC", "EI", "OF", "rho PCI", "rho RPCI", "MAD PCI", "MAD RPCI"]
    compact9 = doc.add_table(rows=len(table9) + 1, cols=len(headers9))
    compact9.style = table_style
    compact9.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(compact9.rows[0].cells, headers9):
        set_cell(cell, header, bold=True, size=6)
    for row_idx, (_, item) in enumerate(table9.iterrows(), start=1):
        values = [item["alpha"], int(item["cycle"]), item["EV"], item["RC"], item["EI"], item["OF"], item["pci_spearman_vs_cycle1"], item["rpci_spearman_vs_cycle1"], item["mean_abs_delta_pci"], item["mean_abs_delta_rpci"]]
        for cell, value in zip(compact9.rows[row_idx].cells, values):
            text = f"{value:.3f}" if isinstance(value, float) else str(value)
            set_cell(cell, text, size=6)
    insert_table_after(table9_title, compact9)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
