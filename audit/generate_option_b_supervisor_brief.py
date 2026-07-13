"""Create a concise supervisor-ready Option B completion brief."""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parents[1]
OUT = WORKSPACE / "output" / "doc" / "AIML_MRMS_Option_B_Supervisor_Brief.docx"


def latest_option_b_dir() -> Path:
    candidates = sorted((ROOT / "audit_artifacts").glob("*_option_b"))
    if not candidates:
        raise FileNotFoundError("No Option B audit package found")
    return candidates[-1]


def style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for name in ["Title", "Heading 1", "Heading 2"]:
        styles[name].font.name = "Aptos Display"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1" if "Light Shading Accent 1" in [s.name for s in doc.styles] else table.style
    table.autofit = True
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, data):
            cell.text = str(text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def main() -> None:
    option_b = latest_option_b_dir()
    matrix = pd.read_csv(option_b / "table8_manuscript_matrix.csv")
    stability = pd.read_csv(option_b / "candidate_table9c_ranking_stability.csv")
    manifest = pd.read_csv(option_b / "configuration_execution_manifest.csv")
    summary = json.loads((option_b / "analysis_manifest.json").read_text(encoding="utf-8"))
    audit_dirs = sorted((ROOT / "audit_artifacts").glob("20*Z"))
    latest_audit = audit_dirs[-1] if audit_dirs else None

    doc = Document()
    style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph("AIML-MRMS Tables 8-9: Option B Completion Brief")
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("For supervisor review - revised weighted-scenario and sensitivity analysis")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Decision requested", level=1)
    doc.add_paragraph(
        "Approve replacement of the withdrawn hard-coded comparison with the attached reproducible weighted-scenario analysis, subject to the revised manuscript scope below. The analysis should be described as sensitivity of PCI/RPCI to explicitly defined criterion-weighting scenarios, not as proof of architecture superiority, closed outcome feedback, or convergence."
    )

    doc.add_heading("Bottom line", level=1)
    doc.add_paragraph(
        "Option B is now computationally reproducible and manuscript-ready as a bounded methodology extension. It preserves a defensible contribution: transparent integration of expert AHP priorities, an aggregate SVM-derived criterion signal, and inequality-aware PCI/RPCI scoring, with parameter and ranking robustness reported. It does not validate an empirically closed feedback loop."
    )

    doc.add_heading("Response to the two supervisory questions", level=1)
    doc.add_heading("1. Are the right things being compared?", level=2)
    doc.add_paragraph(
        "The executed comparisons are now explicitly defined as weighting scenarios. They share the same country input vectors and differ only in criterion-weight source or fixed update count. The implementation does not execute four independent decision architectures."
    )
    add_table(doc, ["Scenario", "Implemented weight source", "Updates", "Feedback/convergence supported?"], [
        [r.scenario, r.implemented_weight_source, str(r.adaptive_updates), "No"]
        for r in manifest.itertuples(index=False)
    ])
    doc.add_heading("2. Why do scores move up or down?", level=2)
    doc.add_paragraph(
        "The package reports the weighted mean, weighted CV penalty, weighted Gini penalty, PCI, and RPCI for every country-scenario pair. It therefore permits a measured explanation of individual score movements. It does not assume that lower adaptive scores are a metric artefact or that higher scores prove a better method."
    )

    doc.add_heading("What has been produced", level=1)
    bullet(doc, "Table 8: 16-country PCI/RPCI matrix for the original unweighted reference, base AHP weights, aggregate SVM signal, one blend, and two blends.")
    bullet(doc, "Table 9: alpha/update-count weight states, score sensitivity, and rank-stability summary.")
    bullet(doc, "Full country-level diagnostic decomposition and exact reproducibility outputs.")
    bullet(doc, "A revised manuscript draft that replaces unsupported claims and incorporates Tables 8 and 9.")

    doc.add_heading("Key robustness result", level=1)
    doc.add_paragraph(
        f"Across the tested grid (alpha = {', '.join(str(x) for x in summary['alpha_values'])}; fixed states = {', '.join(str(x) for x in summary['fixed_iteration_counts'])}), PCI ranking correlations relative to the base-AHP state range from {stability['pci_spearman_vs_cycle1'].min():.3f} to {stability['pci_spearman_vs_cycle1'].max():.3f}; RPCI correlations range from {stability['rpci_spearman_vs_cycle1'].min():.3f} to {stability['rpci_spearman_vs_cycle1'].max():.3f}. This reports robustness transparently without treating ranking stability as a universal performance claim."
    )

    doc.add_heading("Mandatory manuscript changes", level=1)
    add_table(doc, ["Area", "Required position"], [
        ["Methods", "Add weighted PCI/RPCI equations, dimension-weight mapping, original-unweighted reference, alpha grid, fixed update counts, and ranking-stability statistic."],
        ["Results", "Report mixed score changes and use the mean/CV/Gini decomposition for country-specific interpretation."],
        ["Claims", "Remove universal-improvement, benchmark-superiority, executed-feedback, and convergence claims."],
        ["Architecture", "Retain feedback as a proposed operational interface; distinguish it from the fixed-state experiment reported here."],
        ["Limitations", "State the n=16 sample, aggregate SVM signal, fixed update rule, and absence of outcome/TOPSIS/investment feedback."],
    ])

    doc.add_heading("Verification completed", level=1)
    bullet(doc, "Uniform dimension weights reproduce the original unweighted PCI/RPCI formula for all 16 countries.")
    bullet(doc, "Four saved weighted-output files reproduce from the current code and inputs to numerical tolerance.")
    bullet(doc, "The legacy hard-coded gain path is identified and excluded from the revised evidence claim.")
    bullet(doc, "The revision draft has structural checks for the replacement sections and embedded Table 8/9 summaries.")

    doc.add_heading("Remaining author/supervisor decisions", level=1)
    bullet(doc, "Approve the FDI-to-EV criterion mapping used to connect the PCI dimensions to the AHP/SVM criterion weights.")
    bullet(doc, "Approve the bounded wording above as the paper's empirical claim.")
    bullet(doc, "Confirm whether Tables 8 and 9 belong in the main manuscript or whether the full matrices should move to supplementary material with compact main-text summaries.")
    bullet(doc, "Complete a final visual review of the Word tables before submission.")

    doc.add_paragraph()
    source = doc.add_paragraph(
        f"Evidence package: {option_b.relative_to(ROOT)}. Latest reproducibility audit: {latest_audit.relative_to(ROOT) if latest_audit else 'not found'}."
    )
    source.runs[0].italic = True
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
