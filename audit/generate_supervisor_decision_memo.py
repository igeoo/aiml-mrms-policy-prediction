"""Generate the Final Supervisor Decision Memo as DOCX."""

import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path(r"c:\Users\USER\Documents\python_codes\super_project\AIML_MRMS\output\doc")

def main():
    doc = Document()

    # ---- Title ----
    title = doc.add_heading("AIML-MRMS Version C: Final Supervisor Decision Memo", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "Purpose: To obtain final supervisor sign-off on the Version C methodology "
        "and novelty framing, enabling immediate integration into the authoritative manuscript."
    ).runs[0].bold = True

    # ---- Section 1: Peer Reviewer Sign-Off ----
    doc.add_heading("1. Independent Peer Reviewer Sign-Off", level=1)
    doc.add_paragraph(
        "The independent peer reviewer has reviewed the complete Phase A-D remediation package "
        "(including the correlated-error sensitivity scenarios, disaggregated WGI dimensions, "
        "and revised limitations). They have officially signed off on the computational and "
        "scientific rigor of the updates."
    )
    
    doc.add_paragraph("Key quotes from the reviewer's verdict:", style="Body Text").runs[0].italic = True
    quotes = [
        '"All four gaps from my previous check are now closed, and closed properly rather than cosmetically."',
        '"The correlated-error result is the most important one... PCI\'s own uncertainty was understated under the independent-error assumption, and the revision doesn\'t hide that... That\'s the right way to handle a finding that complicates your own results: report it, don\'t bury it."',
        '"This package is ready to move to manuscript integration once you sign off."'
    ]
    for q in quotes:
        p = doc.add_paragraph(q, style="Quote")

    # ---- Section 2: Supervisor Decisions Required ----
    doc.add_heading("2. Decisions Required for Manuscript Integration", level=1)
    doc.add_paragraph(
        "The reviewer endorsed specific options for the remaining open decisions. "
        "Please indicate your approval for the following four items so we can update the manuscript:"
    )

    # Decision 1
    doc.add_heading("Decision 1: The Revised Novelty Statement", level=2)
    doc.add_paragraph("Reviewer recommendation: Approve the fuller formulation.").runs[0].bold = True
    doc.add_paragraph(
        "Proposed text: \"This study contributes a rigor-oriented governance screening protocol "
        "for mineral-investment decision support. It separates country-level governance diagnostics "
        "from project-level economic and operational evaluation, propagates published WGI measurement "
        "uncertainty, evaluates alternative weighting assumptions, conducts leakage-free convergent "
        "validation against a mining-policy benchmark, and transparently reports a null "
        "construct-boundary test against general FDI.\""
    )
    doc.add_paragraph("[   ] APPROVED     [   ] REVISE", style="List Bullet")

    # Decision 2
    doc.add_heading("Decision 2: WGI Limitations Paragraph", level=2)
    doc.add_paragraph("Reviewer recommendation: Approve the explicit World Bank FAQ connection.").runs[0].bold = True
    doc.add_paragraph(
        "Proposed text: \"The World Bank does not publish a single composite across the six WGI "
        "dimensions because such an aggregate would be conceptually broad, correlations partly reflect "
        "shared source inputs, and aggregate margins of error are difficult to construct. The diagnostics "
        "developed here are therefore not presented as official WGI measures or comprehensive representations "
        "of governance. They are author-defined screening summaries...\""
    )
    doc.add_paragraph("[   ] APPROVED     [   ] REVISE", style="List Bullet")

    # Decision 3
    doc.add_heading("Decision 3: Reporting Conservative Correlated Error", level=2)
    doc.add_paragraph("Reviewer recommendation: Approve reporting the widest scenario.").runs[0].bold = True
    doc.add_paragraph(
        "Proposal: In the limitations section, we will explicitly state that under the most conservative "
        "correlated-error scenario tested (ρ=0.75), diagnostic intervals widen by up to 68% compared to "
        "the independent-error baseline, confirming that the baseline understates measurement uncertainty."
    )
    doc.add_paragraph("[   ] APPROVED     [   ] REVISE", style="List Bullet")

    # Decision 4
    doc.add_heading("Decision 4: Renaming PCI/RPCI to GCSI/RGCSI", level=2)
    doc.add_paragraph("Reviewer recommendation: Approve the rename to break historical baggage.").runs[0].bold = True
    doc.add_paragraph(
        "Proposal: Rename 'PCI' (Policy Coherence Index) to 'GCSI' (Governance Coherence Screening Index) "
        "throughout the manuscript to prevent readers or future reviewers from accidentally importing assumptions "
        "from earlier drafts about adaptive feedback or overall superiority."
    )
    doc.add_paragraph("[   ] APPROVED     [   ] REVISE", style="List Bullet")
    
    # ---- Next Steps ----
    doc.add_heading("Next Steps", level=1)
    doc.add_paragraph(
        "Upon your approval of these four items, the implementation team will execute "
        "Phase E: updating the authoritative manuscript document with this text and running "
        "the final automated manuscript-vs-repository consistency audit."
    )

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / "AIML_MRMS_Supervisor_Final_Decision_Memo.docx"
    doc.save(out_path)
    print(f"Memo saved to {out_path}")

if __name__ == "__main__":
    main()
