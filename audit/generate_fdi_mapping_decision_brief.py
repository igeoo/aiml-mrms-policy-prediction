"""Generate a supervisor decision brief for the FDI-to-EV mapping alternatives."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parents[1]
OUTPUT = WORKSPACE / "output" / "doc" / "AIML_MRMS_FDI_to_EV_Mapping_Decision_Brief.docx"


BLUE = RGBColor(31, 78, 121)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = BLUE


def add_title(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("AIML-MRMS: FDI-to-EV Mapping Decision Brief")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Three methodological alternatives for supervisor review")
    run.italic = True
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None, font_size=8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.autofit = True
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)


def add_labelled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(label + ": ")
    r.bold = True
    p.add_run(text)


def add_manuscript_text(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_title(doc)

    doc.add_heading("Purpose and decision requested", level=1)
    doc.add_paragraph(
        "The Option B analysis is reproducible, but the manuscript must explicitly define how the four AHP/SVM criterion weights - Economic Viability (EV), Regulatory Compliance (RC), ESG/Environmental Index (EI), and Operational Feasibility (OF) - become weights over the PCI dimensions. The decision concerns whether normalised FDI inflow should represent EV directly, share the EV weight with governance indicators, or remain outside weighted PCI/RPCI."
    )
    doc.add_paragraph(
        "Approval is requested for one of the three alternatives below before any further recalculation or manuscript insertion. The alternatives are methodological choices, not interchangeable wordings."
    )

    doc.add_heading("Executive comparison", level=1)
    add_table(
        doc,
        ["Version", "What it offers", "Recalculation", "Principal advantage", "Principal risk"],
        [
            ["A. Direct proxy", "Six-dimension PCI with EV assigned to FDI", "None", "Already implemented, transparent and reproducible", "FDI alone may be considered too narrow a proxy for EV"],
            ["B. Weighted split", "Six-dimension PCI with EV divided between FDI and investment-climate indicators", "Full Tables 8-9", "Represents realised investment and enabling governance together", "Overlapping construct weights and a 50/50 split require justification"],
            ["C. Governance-only", "Five-dimension weighted PCI; FDI reported separately", "Full Tables 8-9 plus FDI check", "Avoids treating FDI as EV", "Changes the PCI construct and removes EV from weighted scoring"],
        ],
        font_size=7,
    )

    doc.add_heading("Current verified implementation", level=1)
    doc.add_paragraph(
        "The existing Option B code implements Version A. For a criterion-weight vector w = (wEV, wRC, wEI, wOF), the current dimension weights are:"
    )
    add_table(
        doc,
        ["PCI dimension", "Current assignment", "Source"],
        [
            ["Normalised FDI inflow", "wEV", "UNCTAD FDI, 2022 value"],
            ["Regulatory Quality", "wRC / 2", "World Bank WGI 2022"],
            ["Rule of Law", "wRC / 2", "World Bank WGI 2022"],
            ["Control of Corruption", "wEI", "World Bank WGI 2022"],
            ["Government Effectiveness", "wOF / 2", "World Bank WGI 2022"],
            ["Political Stability", "wOF / 2", "World Bank WGI 2022"],
        ],
    )

    doc.add_page_break()
    doc.add_heading("Version A - Direct FDI proxy for EV", level=1)
    add_labelled_paragraph(doc, "Methodological definition", "FDI is the sixth PCI dimension and receives the full EV criterion weight. RC and OF are divided equally across their two mapped WGI dimensions; EI maps to Control of Corruption.")
    add_labelled_paragraph(doc, "What this version offers", "A simple, fully specified link between the existing four-criterion weight vector and all six PCI dimensions. It preserves the current Option B tables and the manuscript's use of UNCTAD FDI within PCI/RPCI.")
    add_labelled_paragraph(doc, "Implementation implication", "No Table 8 or Table 9 recalculation is required. The missing mapping paragraph and mapping table are added to Section 6.5, and the proxy assumption is acknowledged in the limitations.")
    add_labelled_paragraph(doc, "Reviewer-facing implication", "The calculation is easy to reproduce, but reviewers may question whether a single observed capital-flow indicator captures the broader construct of economic viability. The manuscript should describe FDI as an operational proxy, not as a complete measurement of EV.")
    doc.add_heading("Mapping", level=2)
    add_table(
        doc,
        ["Criterion", "Assigned PCI dimension(s)", "Dimension-weight rule"],
        [
            ["EV", "FDI inflow", "FDI = wEV"],
            ["RC", "Regulatory Quality; Rule of Law", "Each = wRC / 2"],
            ["EI", "Control of Corruption", "Control of Corruption = wEI"],
            ["OF", "Government Effectiveness; Political Stability", "Each = wOF / 2"],
        ],
    )
    doc.add_heading("Proposed Section 6.5 paragraph", level=2)
    add_manuscript_text(
        doc,
        "The six PCI dimensions comprise five World Bank WGI indicators - Rule of Law, Regulatory Quality, Government Effectiveness, Control of Corruption, and Political Stability - plus normalised UNCTAD FDI inflow as the investment-flow dimension. Within the four-criterion AHP/SVM vector, FDI inflow is used as the operational proxy for Economic Viability (EV), because realised inbound capital provides an observable jurisdiction-level signal of perceived investment viability. The EV weight is assigned to FDI; RC is divided equally between Regulatory Quality and Rule of Law; EI is assigned to Control of Corruption; and OF is divided equally between Government Effectiveness and Political Stability. This proxy mapping is an explicit modelling assumption and should not be interpreted as equating FDI with the full conceptual scope of economic viability."
    )

    doc.add_page_break()
    doc.add_heading("Version B - Split EV between FDI and investment-climate indicators", level=1)
    add_labelled_paragraph(doc, "Methodological definition", "Half of EV is assigned to FDI. The remaining half is divided equally between Regulatory Quality and Government Effectiveness, in addition to their existing RC and OF contributions.")
    add_labelled_paragraph(doc, "What this version offers", "A broader operationalisation of EV that reflects both realised capital flow and two governance conditions associated with investment activity. It reduces dependence on FDI as a single proxy.")
    add_labelled_paragraph(doc, "Implementation implication", "All weighted PCI/RPCI values, score decompositions, and alpha/state sensitivity results must be recalculated. Table 8, Table 9, the reproducibility supplement, and the manuscript results narrative must be replaced and re-audited.")
    add_labelled_paragraph(doc, "Reviewer-facing implication", "The approach may be conceptually attractive, but Regulatory Quality and Government Effectiveness receive contributions from more than one criterion. The 50/50 split and overlapping assignments require a theoretical or expert-elicitation basis; otherwise they may appear arbitrary or double-counted.")
    doc.add_heading("Mapping", level=2)
    add_table(
        doc,
        ["PCI dimension", "Dimension-weight rule"],
        [
            ["FDI inflow", "0.50 wEV"],
            ["Regulatory Quality", "0.50 wRC + 0.25 wEV"],
            ["Rule of Law", "0.50 wRC"],
            ["Control of Corruption", "wEI"],
            ["Government Effectiveness", "0.50 wOF + 0.25 wEV"],
            ["Political Stability", "0.50 wOF"],
        ],
    )
    doc.add_heading("Proposed Section 6.5 paragraph", level=2)
    add_manuscript_text(
        doc,
        "The six PCI dimensions comprise the five WGI governance indicators plus normalised UNCTAD FDI inflow. Economic Viability is operationalised as a combination of realised capital flow and governance conditions that support investment. Accordingly, 50% of the EV criterion weight is assigned to FDI, while the remaining 50% is divided equally between Regulatory Quality and Government Effectiveness. These EV-derived contributions are added to the dimensions' existing RC- and OF-derived allocations. Rule of Law receives half of RC, Control of Corruption receives EI, and Political Stability receives half of OF. The resulting six dimension weights sum to one. The 50/50 split is a substantive modelling assumption and requires explicit author approval and theoretical justification."
    )
    doc.add_heading("Required downstream work", level=2)
    add_bullet(doc, "Implement the new dimension-weight function and unit tests confirming weights sum to one.")
    add_bullet(doc, "Regenerate Table 8, all decompositions, and Table 9 sensitivity/ranking results.")
    add_bullet(doc, "Repeat numerical reproducibility and manuscript-to-table consistency checks.")
    add_bullet(doc, "Revise Section 6.5, limitations, Figure 4 caption, supplement, and results interpretation.")

    doc.add_page_break()
    doc.add_heading("Version C - Governance-only weighted PCI with FDI reported separately", level=1)
    add_labelled_paragraph(doc, "Clarification applied", "The task brief did not specify what happens to EV after FDI is removed. For this version to be executable, EV is excluded from weighted PCI scoring. At each weight state, RC, EI and OF are renormalised to sum to one before they are mapped across the five WGI dimensions. FDI is retained as a separate descriptive robustness indicator.")
    add_labelled_paragraph(doc, "What this version offers", "A governance-coherence index that avoids treating FDI as a proxy for economic viability. It creates a clean separation between governance alignment and observed investment flow.")
    add_labelled_paragraph(doc, "Implementation implication", "This is not a small mapping edit. It changes weighted PCI/RPCI from a six-dimension, four-criterion construct into a five-dimension, three-criterion governance score. All Table 8 and Table 9 values must be recalculated, and a separate FDI robustness table or analysis must be introduced.")
    add_labelled_paragraph(doc, "Reviewer-facing implication", "It removes the most contestable proxy assumption, but the manuscript can no longer imply that weighted PCI directly integrates economic viability or FDI. The distinction between the four-criterion architecture and the three-criterion PCI scoring projection must be stated clearly.")
    doc.add_heading("Mapping", level=2)
    add_table(
        doc,
        ["Criterion", "Treatment in weighted PCI", "Assigned WGI dimension(s)"],
        [
            ["EV", "Excluded from scoring", "None; FDI reported separately"],
            ["RC", "Renormalised RC / (RC + EI + OF)", "Regulatory Quality; Rule of Law, equal split"],
            ["EI", "Renormalised EI / (RC + EI + OF)", "Control of Corruption"],
            ["OF", "Renormalised OF / (RC + EI + OF)", "Government Effectiveness; Political Stability, equal split"],
        ],
    )
    doc.add_heading("Proposed Section 6.5 paragraph", level=2)
    add_manuscript_text(
        doc,
        "The weighted PCI/RPCI formulation is restricted to the five WGI governance dimensions. Economic Viability and FDI are excluded from the weighted governance score. For each weight state, the RC, EI and OF criterion weights are renormalised to sum to one; RC is divided equally between Regulatory Quality and Rule of Law, EI is assigned to Control of Corruption, and OF is divided equally between Government Effectiveness and Political Stability. Normalised UNCTAD FDI inflow is reported separately as a descriptive robustness indicator, permitting comparison between governance-weighted PCI/RPCI patterns and observed investment flow without combining them in a single composite score. This version therefore measures governance alignment rather than the broader four-criterion AIML-MRMS construct."
    )
    doc.add_heading("Required downstream work", level=2)
    add_bullet(doc, "Define and test the RC/EI/OF renormalisation at every scenario and alpha/state combination.")
    add_bullet(doc, "Regenerate Table 8 and Table 9 using five WGI dimensions and the governance-only reference score.")
    add_bullet(doc, "Add a separate FDI robustness analysis and define its statistic before running it.")
    add_bullet(doc, "Revise the PCI definition, Figure 4, data description, abstract/results claims, limitations, and supplement.")

    doc.add_page_break()
    doc.add_heading("Decision implications", level=1)
    add_table(
        doc,
        ["Decision consideration", "Version A", "Version B", "Version C"],
        [
            ["Preserves current Option B outputs", "Yes", "No", "No"],
            ["Keeps FDI inside PCI/RPCI", "Yes, full EV proxy", "Yes, partial EV proxy", "No"],
            ["Avoids single-indicator EV proxy", "No", "Partly", "Yes"],
            ["Preserves four-criterion scoring", "Yes", "Yes", "No; PCI scoring uses RC/EI/OF"],
            ["New modelling assumptions", "FDI = operational EV proxy", "50/50 EV split and overlapping assignments", "EV exclusion and governance-only reinterpretation"],
            ["Implementation effort", "Low", "Medium-high", "High"],
            ["Submission risk", "Proxy may be challenged", "Split may appear arbitrary", "Construct change may require broad rewriting"],
        ],
        font_size=7,
    )

    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph(
        "For the current revision cycle, Version A is recommended because it is already implemented, reproducible, and consistent with the completed Option B tables. Its limitation can be managed by explicitly identifying FDI as an operational proxy and acknowledging that it does not exhaust the EV construct. Version B should be selected only if the supervisor can support the 50/50 split with theory or expert elicitation. Version C is appropriate if conceptual separation between governance alignment and investment flow is more important than preserving the present PCI construct, but it requires the broadest recalculation and manuscript revision."
    )

    doc.add_heading("Approval requested", level=1)
    add_bullet(doc, "Select Version A, B or C as the approved dimension-mapping method.")
    add_bullet(doc, "If Version B is selected, confirm or replace the proposed 50/50 EV split.")
    add_bullet(doc, "If Version C is selected, approve exclusion of EV from weighted PCI and the RC/EI/OF renormalisation rule.")
    add_bullet(doc, "Confirm whether the mapping table belongs in Section 6.5 or in the main text with full detail repeated in the supplement.")

    note = doc.add_paragraph()
    note.add_run("Prepared from the verified Option B implementation and the FDI-variation task brief. No Tables 8-9 have been recalculated pending methodological approval.").italic = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
