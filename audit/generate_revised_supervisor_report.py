"""Generate the revised Version C supervisor report.

Assembles all Phase A-C outputs into a single DOCX report for supervisor review.
Does NOT silently replace the interim report.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

import pandas as pd

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("python-docx required: pip install python-docx")
    sys.exit(1)


ROOT = Path(r"c:\Users\USER\Documents\python_codes\super_project\AIML_MRMS"
            r"\aiml_mrms_github_package\aiml_mrms_github_package")
TABLES = ROOT / "results" / "tables"
CORR_DIR = ROOT / "audit_artifacts" / "20260717_correlated_error_sensitivity_run1"
ROBUST_DIR = ROOT / "audit_artifacts" / "20260717_six_dimension_robustness_final"
OUTPUT_DIR = Path(r"c:\Users\USER\Documents\python_codes\super_project\AIML_MRMS"
                  r"\output\doc")


def add_table_from_df(doc, df, title=None, font_size=8):
    """Add a formatted table from a DataFrame."""
    if title:
        doc.add_heading(title, level=3)
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(font_size)
    for i, (_, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            if isinstance(val, float):
                cell.text = f"{val:.4f}"
            else:
                cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def main():
    doc = Document()

    # Title
    title = doc.add_heading("AIML-MRMS Version C: Revised Supervisor Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "This report supersedes the interim Version C report. It incorporates "
        "the required correlated-error sensitivity analysis, disaggregated WGI "
        "dimension reporting, revised novelty framing, and expanded prior-art "
        "comparison. It is submitted for supervisor/author review before "
        "manuscript integration."
    )
    doc.add_paragraph(
        "Status: AWAITING SUPERVISOR APPROVAL on novelty statement and "
        "limitations paragraph before manuscript integration."
    ).runs[0].bold = True

    # === Section 1: Executive Summary ===
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Version C presents AIML-MRMS as a governance-aware screening protocol "
        "for mineral-investment decision support. It does not claim superiority "
        "over alternative architectures, causality, convergence, or adaptive "
        "feedback. The former cross-configuration PCI comparison (Tables 8/9) "
        "remains withdrawn."
    )
    doc.add_paragraph("Key findings from this revision:")
    bullets = [
        "Correlated WGI measurement-error sensitivity: PCI/RPCI rank ordering "
        "is perfectly stable (Spearman ρ = 1.000) across all equicorrelation "
        "scenarios (ρ_error = 0.00 to 0.75). Four countries show material "
        "rank-span widening at high correlation.",
        "PCI intervals widen by up to 68% at ρ_error = 0.75, confirming the "
        "supervisor's concern that independent-error sampling may understate "
        "uncertainty. The widest-scenario intervals are now reported.",
        "All six WGI dimensions are reported separately (Table 8A) alongside "
        "the composite screening diagnostics (Table 8B).",
        "The null general-FDI result is retained and correctly labelled as an "
        "audit-introduced construct-boundary test, not a pre-specified "
        "falsification.",
        "The novelty claim is narrowed to a rigor-and-validation contribution.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # === Section 2: Correlated-Error Sensitivity Results ===
    doc.add_heading("2. Correlated WGI Measurement-Error Sensitivity", level=1)
    doc.add_paragraph(
        "The independent-error baseline sampled six WGI dimension errors "
        "independently. Given that WGI dimensions share underlying data sources "
        "(reflected in VIF = 23.12), positive error correlation is plausible. "
        "This section reports equicorrelation sensitivity scenarios."
    )
    doc.add_paragraph(
        "Assumption: These are sensitivity scenarios, not estimates of the true "
        "error covariance. Observed WGI score correlation is not measurement-error "
        "correlation."
    ).runs[0].italic = True

    # Load and display scenario summary
    summary = pd.read_csv(CORR_DIR / "correlated_error_scenario_summary.csv")
    summary_display = summary[[
        "rho", "mean_pci_interval_width", "mean_pci_width_ratio_vs_rho0",
        "max_pci_width_ratio_vs_rho0", "pci_rank_spearman_vs_rho0",
    ]].rename(columns={
        "rho": "ρ_error",
        "mean_pci_interval_width": "Mean PCI 95% Width",
        "mean_pci_width_ratio_vs_rho0": "Mean Width Ratio vs ρ=0",
        "max_pci_width_ratio_vs_rho0": "Max Width Ratio vs ρ=0",
        "pci_rank_spearman_vs_rho0": "PCI Rank Spearman vs ρ=0",
    })
    add_table_from_df(doc, summary_display, "Table 2.1: Correlated-Error Scenario Summary")

    # Material rank changes
    stability = pd.read_csv(CORR_DIR / "correlated_error_rank_stability.csv")
    material = stability[stability["material_change"]]
    if len(material) > 0:
        doc.add_heading("Countries with Material Rank-Span Changes (≥2 positions)", level=3)
        material_display = material[[
            "iso3", "rho", "pci_rank_span_rho0", "pci_rank_span_correlated",
            "pci_rank_span_change"
        ]].rename(columns={
            "iso3": "Country",
            "rho": "ρ_error",
            "pci_rank_span_rho0": "PCI Rank Span (ρ=0)",
            "pci_rank_span_correlated": "PCI Rank Span (correlated)",
            "pci_rank_span_change": "Change",
        })
        add_table_from_df(doc, material_display)

    doc.add_paragraph(
        "Interpretation: Positive error correlation widens PCI/RPCI intervals "
        "as expected, but does not alter the broad rank ordering. The widest-"
        "scenario (ρ_error = 0.75) intervals should be reported in the "
        "limitations section. Fine-grained league-table claims should not be "
        "made between countries with overlapping intervals."
    )

    # === Section 3: Disaggregated WGI Dimensions ===
    doc.add_heading("3. Six-Dimension WGI Profile (Table 8A)", level=1)
    doc.add_paragraph(
        "All six official WGI dimensions are reported separately so the reader "
        "can assess the dimension-level evidence rather than relying solely on "
        "the author-defined composite."
    )

    table8a = pd.read_csv(TABLES / "table8a_wgi_six_dimensions_2024.csv")
    # Display a compact version
    compact_cols = ["country", "iso3"]
    dim_labels = [
        "Control of Corruption", "Government Effectiveness",
        "Political Stability", "Regulatory Quality",
        "Rule of Law", "Voice and Accountability",
    ]
    for d in dim_labels:
        compact_cols.append(d)
    compact_cols.extend(["pci", "rpci", "pci_rank"])
    compact = table8a[compact_cols].copy()
    for d in dim_labels:
        compact[d] = compact[d].round(3)
    compact["pci"] = compact["pci"].round(4)
    compact["rpci"] = compact["rpci"].round(4)
    add_table_from_df(doc, compact, "Table 8A: WGI Six-Dimension Profile, 2024", font_size=7)

    # === Section 4: Governance Coherence Diagnostics ===
    doc.add_heading("4. Governance Coherence Diagnostics (Table 8B)", level=1)
    table8b = pd.read_csv(TABLES / "table8b_governance_coherence_diagnostics_2024.csv")
    add_table_from_df(doc, table8b, "Table 8B: PCI/RPCI Screening Diagnostics with Measurement Uncertainty")
    doc.add_paragraph(
        "PCI and RPCI are author-defined governance-coherence screening "
        "diagnostics. They are not official World Bank measures and should not "
        "be presented as definitive overall governance scores."
    ).runs[0].italic = True

    # === Section 5: Validation and Robustness (Table 9) ===
    doc.add_heading("5. Validation and Robustness Evidence (Table 9)", level=1)
    table9 = pd.read_csv(TABLES / "table9_version_c_validation_robustness.csv")
    add_table_from_df(doc, table9, "Table 9: Validation and Robustness Summary")

    # === Section 6: Novelty Evidence Matrix ===
    doc.add_heading("6. Novelty Evidence Matrix", level=1)
    doc.add_paragraph(
        "The following matrix compares Version C's methodological features "
        "against identified prior art in extractive-sector country-risk MCDM."
    )
    matrix = pd.read_csv(TABLES / "novelty_evidence_matrix.csv")
    add_table_from_df(doc, matrix, font_size=7)

    # === Section 7: Revised Novelty Statement ===
    doc.add_heading("7. Revised Novelty Statement (FOR SUPERVISOR APPROVAL)", level=1)
    doc.add_paragraph("Recommended working statement:").bold = True
    doc.add_paragraph(
        "This study contributes a rigor-oriented governance screening protocol "
        "for mineral-investment decision support. It separates country-level "
        "governance diagnostics from project-level economic and operational "
        "evaluation, propagates published WGI measurement uncertainty, evaluates "
        "alternative weighting assumptions, conducts leakage-free convergent "
        "validation against a mining-policy benchmark, and transparently reports "
        "a null construct-boundary test against general FDI."
    )
    doc.add_paragraph("Alternative compact formulation:").bold = True
    doc.add_paragraph(
        "An author-defined mineral-investment governance-coherence screening "
        "layer distinguished by construct separation, WGI uncertainty propagation, "
        "leakage-free mining-policy validation, explicit null-result reporting, "
        "and reproducible weighting and temporal stress tests."
    )
    doc.add_paragraph("This study does NOT claim:")
    prohibited = [
        "A novel governance construct.",
        "A novel uncertainty or Monte Carlo method.",
        "A novel AHP-TOPSIS architecture.",
        "A mining-specific dataset (WGI inputs are general governance data).",
        "A pre-specified FDI falsification test.",
        "A validated adaptive feedback loop.",
        "Superiority over alternative configurations.",
        "Causality between governance and investment outcomes.",
    ]
    for p in prohibited:
        doc.add_paragraph(p, style="List Bullet")

    # === Section 8: WGI Limitations Paragraph ===
    doc.add_heading("8. WGI Limitations Paragraph (FOR SUPERVISOR APPROVAL)", level=1)
    doc.add_paragraph(
        "The World Bank does not publish a single composite across the six WGI "
        "dimensions because such an aggregate would be conceptually broad, "
        "correlations partly reflect shared source inputs, and aggregate margins "
        "of error are difficult to construct. The diagnostics developed here are "
        "therefore not presented as official WGI measures or comprehensive "
        "representations of governance. They are author-defined screening "
        "summaries for examining the level and balance of governance dimensions "
        "within the study context. All six dimensions remain reported separately, "
        "and any composite results are interpreted alongside dimension profiles, "
        "measurement uncertainty, correlated-error sensitivity, and alternative-"
        "weight analyses."
    )

    # === Section 9: Prior Art References ===
    doc.add_heading("9. Required Prior-Art Additions", level=1)
    refs = [
        "Tafur, Lilford, and Aguilera (2022). Petroleum investment-risk ranking "
        "in South America using AHP and TOPSIS. https://doi.org/10.1007/s43546-022-00221-6",
        "Tang et al. (2023). Oil and gas investment-risk assessment across 76 "
        "Belt and Road countries using AHP/entropy and TOPSIS-GRA. "
        "https://doi.org/10.1016/j.petsci.2023.10.009",
        "Li et al. (2020). Investment risk and natural-resource potential across "
        "63 Belt and Road countries using entropy-TOPSIS. "
        "https://doi.org/10.1016/j.scitotenv.2020.137981",
        "Mining IQ / MineHutte World Risk Insights methodology. "
        "https://www.mining-iq.com/focus-risk-methodology (commercial/practitioner, "
        "not necessarily peer-reviewed).",
        "World Bank WGI FAQ on composite scores. "
        "https://www.worldbank.org/en/publication/worldwide-governance-indicators/"
        "frequently-asked-questions",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Number")

    # === Section 10: Remaining Author Decisions ===
    doc.add_heading("10. Remaining Author/Supervisor Decisions", level=1)
    decisions = [
        "Approve or revise the novelty statement in Section 7.",
        "Approve or revise the WGI limitations paragraph in Section 8.",
        "Decide whether to rename PCI/RPCI to GCSI/RGCSI to reduce confusion "
        "with historical claims.",
        "Identify the authoritative manuscript file for Phase E updates.",
        "Confirm that the correlated-error sensitivity at ρ=0.75 should be "
        "reported as the conservative scenario in the limitations section.",
    ]
    for d in decisions:
        doc.add_paragraph(d, style="List Number")

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / "AIML_MRMS_Version_C_Revised_Supervisor_Report.docx"
    doc.save(output_path)
    print(f"Revised supervisor report saved to {output_path}")


if __name__ == "__main__":
    main()
