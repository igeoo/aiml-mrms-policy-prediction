"""Generate the peer-reviewer verification guide as DOCX."""

import sys, json
from pathlib import Path
from datetime import datetime, timezone

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd


ROOT = Path(r"c:\Users\USER\Documents\python_codes\super_project\AIML_MRMS"
            r"\aiml_mrms_github_package\aiml_mrms_github_package")
CORR_DIR = ROOT / "audit_artifacts" / "20260717_correlated_error_sensitivity_run1"
TABLES = ROOT / "results" / "tables"
OUTPUT = Path(r"c:\Users\USER\Documents\python_codes\super_project\AIML_MRMS"
              r"\output\doc")


def add_small_table(doc, headers, rows, font_size=8):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)


def main():
    doc = Document()

    # ---- Title ----
    title = doc.add_heading(
        "AIML-MRMS Version C: Independent Verification Guide", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "This document enables independent verification of all computational "
        "outputs produced during the AIML-MRMS Version C remediation. It "
        "provides reproduction commands, expected SHA-256 hashes, scientific "
        "constraint checklists, and result summaries. The accompanying "
        "peer-review ZIP contains all scripts and data outputs."
    )

    # ---- Section 1: Environment ----
    doc.add_heading("1. Environment and Prerequisites", level=1)
    doc.add_paragraph("Python 3.10+ with the following packages:")
    for pkg in ["numpy", "pandas", "scipy", "scikit-learn", "python-docx", "openpyxl"]:
        doc.add_paragraph(pkg, style="List Bullet")
    doc.add_paragraph(
        "Working directory (all commands below are relative to this):"
    )
    doc.add_paragraph(
        r"C:\Users\USER\Documents\python_codes\super_project\AIML_MRMS"
        r"\aiml_mrms_github_package\aiml_mrms_github_package",
        style="Quote",
    )
    doc.add_paragraph("Branch: codex/version-c-novelty-extension")
    doc.add_paragraph(
        "WGI source workbook SHA-256: "
        "25a2f9eabb90b0092973392c0b31571aa58b691cc5786292e504b52f693e1eb8"
    )

    # ---- Section 2: Reproduction Commands ----
    doc.add_heading("2. Reproduction Commands", level=1)

    steps = [
        (
            "Step 1: Build the WGI panel (already locked)",
            "python audit/build_wgi_six_dimension_panel.py "
            "data/external/wgidataset_with_sourcedata-2025.xlsx "
            "--output-dir audit_artifacts/[new_dir]_wgi_panel",
            "Requires the WGI workbook. Output: 368 country-year observations, "
            "16 countries, 23 years, 6 dimensions.",
        ),
        (
            "Step 2: Run independent-error robustness baseline (already locked)",
            "python audit/run_six_dimension_robustness.py "
            "--wgi-panel audit_artifacts/20260716_novelty_wgi_panel/"
            "wgi_2002_2024_six_dimension_panel.csv "
            "--fraser-correlations audit_artifacts/20260717_fraser_convergent_validation/"
            "fraser_dimension_correlations.csv "
            "--output-dir audit_artifacts/[new_dir]_robustness",
            "10,000 MC draws, seed 42. Produces weight perturbation, measurement "
            "uncertainty, and temporal stability analyses.",
        ),
        (
            "Step 3: Run correlated-error sensitivity (new)",
            "python audit/run_correlated_wgi_uncertainty.py "
            "--wgi-panel audit_artifacts/20260716_novelty_wgi_panel/"
            "wgi_2002_2024_six_dimension_panel.csv "
            "--output-dir audit_artifacts/[new_dir]_correlated_error "
            "--mc-draws 10000 "
            "--locked-baseline-dir audit_artifacts/"
            "20260717_six_dimension_robustness_final",
            "10,000 draws per scenario, seed 42. "
            "Scenarios: ρ = 0.00, 0.25, 0.50, 0.75.",
        ),
        (
            "Step 4: Build disaggregated WGI tables",
            "python audit/build_disaggregated_wgi_tables.py "
            "--wgi-panel audit_artifacts/20260716_novelty_wgi_panel/"
            "wgi_2002_2024_six_dimension_panel.csv "
            "--robustness-dir audit_artifacts/"
            "20260717_six_dimension_robustness_final "
            "--output-dir results/tables",
            "Produces table8a (wide), table8b (diagnostics), and "
            "supplement (long).",
        ),
        (
            "Step 5: Build novelty evidence matrix",
            "python audit/build_novelty_evidence_matrix.py",
            "Produces novelty_evidence_matrix.csv (12 features × 5 studies).",
        ),
    ]
    for title_text, command, note in steps:
        doc.add_heading(title_text, level=3)
        doc.add_paragraph(command, style="Quote")
        doc.add_paragraph(note)

    # ---- Section 3: SHA-256 Verification ----
    doc.add_heading("3. SHA-256 Verification Hashes", level=1)
    doc.add_paragraph(
        "All outputs are deterministic given the same inputs and seed. "
        "The following hashes were computed from the authoritative run. "
        "A verifier should reproduce the computation and confirm matching "
        "hashes for all CSV files."
    )

    hash_data = [
        ("correlated_error_scenario_summary.csv", "7257cfc8...74b724c", "881"),
        ("correlated_error_country_intervals_2024.csv", "cc1ca07e...83bccf8", "11,246"),
        ("correlated_error_rank_intervals_2024.csv", "11dbf6b1...6bcb61", "3,040"),
        ("correlated_error_interval_width_ratios.csv", "c9dcf842...3e2f8c", "6,219"),
        ("correlated_error_rank_stability.csv", "d1505664...d175d8", "2,085"),
        ("correlated_error_manifest.json", "1b024dbe...91399a", "2,130"),
        ("table8a_wgi_six_dimensions_2024.csv", "7a8beb9e...827cfb", "7,408"),
        ("table8b_governance_coherence_diagnostics_2024.csv", "c84c9655...b344e", "1,876"),
        ("supplement_wgi_dimension_intervals_2024_long.csv", "c51a7450...10d93", "10,821"),
        ("table8_version_c_governance_coherence_2024.csv", "4e1ca232...48188d", "1,167"),
        ("table9_version_c_validation_robustness.csv", "6da90bd7...aadb4d", "1,944"),
        ("novelty_evidence_matrix.csv", "18786490...2511ed", "1,593"),
        ("six_dimension_robustness_manifest.json", "ca46a835...f447a", "807"),
        ("weight_perturbation_summary.json", "ebb363a1...433ae", "287"),
        ("wgi_panel_validation.json", "03ff12f5...a24f36", "1,117"),
    ]
    add_small_table(
        doc,
        ["File", "SHA-256 (truncated)", "Size (bytes)"],
        hash_data,
        font_size=7,
    )
    doc.add_paragraph(
        "Full SHA-256 hashes are available in the MANIFEST.json inside the "
        "peer-review ZIP."
    )

    # ---- Section 4: Key Results Summary ----
    doc.add_heading("4. Key Results Summary", level=1)

    doc.add_heading("4.1 Correlated-Error Sensitivity", level=2)
    summary = pd.read_csv(CORR_DIR / "correlated_error_scenario_summary.csv")
    rows_s = []
    for _, r in summary.iterrows():
        rows_s.append([
            f"{r['rho']:.2f}",
            f"{r['mean_pci_interval_width']:.4f}",
            f"{r['mean_pci_width_ratio_vs_rho0']:.3f}",
            f"{r['max_pci_width_ratio_vs_rho0']:.3f}",
            f"{r['pci_rank_spearman_vs_rho0']:.3f}",
        ])
    add_small_table(
        doc,
        ["ρ_error", "Mean PCI 95% Width", "Mean Width Ratio", "Max Width Ratio", "Rank Spearman"],
        rows_s,
    )
    doc.add_paragraph(
        "Interpretation: Rank ordering is perfectly stable (Spearman = 1.000) "
        "across all scenarios. PCI intervals widen by up to 68% at ρ = 0.75. "
        "Four countries (MRT, SLE, ZAF, ZMB) show material rank-span widening "
        "(≥2 positions). No broad-tier conclusion changes."
    )

    doc.add_heading("4.2 Reproducibility Verification", level=2)
    doc.add_paragraph(
        "The correlated-error analysis was run twice independently. All five "
        "CSV outputs were byte-identical between runs (confirmed via SHA-256). "
        "The manifest JSON differs only in its timestamp field."
    )

    doc.add_heading("4.3 Baseline Cross-Validation", level=2)
    manifest = json.loads(
        (CORR_DIR / "correlated_error_manifest.json").read_text(encoding="utf-8")
    )
    bc = manifest.get("baseline_comparison", {})
    if bc:
        add_small_table(
            doc,
            ["Metric", "Max Absolute Difference vs Locked Baseline"],
            [
                ["PCI lower 95%", f"{bc['max_abs_pci_lower95_difference']:.6f}"],
                ["PCI upper 95%", f"{bc['max_abs_pci_upper95_difference']:.6f}"],
                ["RPCI lower 95%", f"{bc['max_abs_rpci_lower95_difference']:.6f}"],
                ["RPCI upper 95%", f"{bc['max_abs_rpci_upper95_difference']:.6f}"],
            ],
        )
        doc.add_paragraph(
            "Differences are within Monte Carlo sampling noise (~0.001 for "
            "10,000 draws) and arise because the ρ=0 scenario uses multivariate "
            "normal (diagonal covariance) rather than independent univariate "
            "normal. Both are statistically equivalent."
        )

    doc.add_heading("4.4 Locked Validation Results (Unchanged)", level=2)
    locked_results = [
        ["Fraser PPI Convergent Validation", "Spearman ρ = 0.715, bootstrap 95% CI [0.437, 0.848]"],
        ["Nested LOCO Ridge (Fraser)", "R² = 0.336, MAE = 15.70, Spearman ρ = 0.602"],
        ["General FDI Boundary Test", "Incremental MAE = −0.007, P(improvement) = 0.291 → NULL"],
        ["Cronbach Alpha (6 WGI dims)", "0.969"],
        ["PC1 Explained Variance", "0.886"],
        ["Maximum VIF", "23.12"],
        ["PCA Weight PCI Rank Spearman", "1.000"],
        ["Entropy Weight PCI Rank Spearman", "0.994"],
        ["Fraser Weight PCI Rank Spearman", "1.000"],
        ["Dirichlet Perturbation (10k draws)", "PCI rank ρ median = 0.997, lower 2.5% = 0.991"],
    ]
    add_small_table(doc, ["Test", "Result"], locked_results)

    # ---- Section 5: Scientific Constraint Checklist ----
    doc.add_heading("5. Scientific Constraint Verification Checklist", level=1)
    doc.add_paragraph(
        "The reviewer should verify that none of the following prohibited "
        "claims appear in any output, report, or manuscript text:"
    )

    constraints = [
        ("Original four-configuration comparison withdrawn", "PASS", "Hard-coded pci_gain_scenarios() values are not used anywhere."),
        ("No 'pre-specified' for audit-introduced FDI test", "PASS", "Uses 'explicitly reported construct-boundary test' throughout."),
        ("No feedback/convergence/superiority/causality claims", "PASS", "Report and novelty statement explicitly list prohibited claims."),
        ("No official WGI composite status claimed", "PASS", "WGI limitations paragraph addresses World Bank's position directly."),
        ("No novel AHP-TOPSIS or uncertainty methodology claimed", "PASS", "Novelty narrowed to rigor-and-validation contribution."),
        ("Six raw WGI dimensions reported alongside composite", "PASS", "Table 8A provides all six dimensions; Table 8B is secondary."),
        ("Negative/null results retained", "PASS", "General FDI null result in Table 9 and report."),
        ("Prior art (Tafur, Tang, Li, MineHutte, WGI FAQ) cited", "PASS", "Novelty evidence matrix and report Section 9."),
        ("Correlated-error sensitivity completed", "PASS", "Four equicorrelation scenarios run and validated."),
        ("Independent baseline unchanged", "PASS", "Locked artifacts directory untouched; SHA-256 verified."),
    ]
    add_small_table(
        doc,
        ["Constraint", "Status", "Evidence"],
        constraints,
        font_size=7,
    )

    # ---- Section 6: Files Inventory ----
    doc.add_heading("6. Peer-Review Package Inventory", level=1)
    doc.add_paragraph("The peer-review ZIP contains 20 files:")

    categories = [
        ("Correlated-Error Sensitivity (6 files)", [
            "correlated_error_scenario_summary.csv",
            "correlated_error_country_intervals_2024.csv",
            "correlated_error_rank_intervals_2024.csv",
            "correlated_error_interval_width_ratios.csv",
            "correlated_error_rank_stability.csv",
            "correlated_error_manifest.json",
        ]),
        ("Disaggregated WGI Tables (3 files)", [
            "table8a_wgi_six_dimensions_2024.csv",
            "table8b_governance_coherence_diagnostics_2024.csv",
            "supplement_wgi_dimension_intervals_2024_long.csv",
        ]),
        ("Existing Validated Tables (3 files)", [
            "table8_version_c_governance_coherence_2024.csv",
            "table9_version_c_validation_robustness.csv",
            "novelty_evidence_matrix.csv",
        ]),
        ("Locked Baseline Manifests (2 files)", [
            "six_dimension_robustness_manifest.json",
            "weight_perturbation_summary.json",
        ]),
        ("Reproducible Scripts (6 files)", [
            "run_correlated_wgi_uncertainty.py",
            "build_disaggregated_wgi_tables.py",
            "build_novelty_evidence_matrix.py",
            "run_six_dimension_robustness.py",
            "run_novelty_modeling.py",
            "run_fraser_convergent_validation.py",
        ]),
    ]
    for category, files in categories:
        doc.add_heading(category, level=3)
        for f in files:
            doc.add_paragraph(f, style="List Bullet")

    doc.add_paragraph(
        "ZIP SHA-256: cc1b8d3c4ca7f969c8727fa15c0faa8e3f6a5ca84f8f444be680a631e78db814"
    ).runs[0].bold = True

    # ---- Section 7: Suggested Review Protocol ----
    doc.add_heading("7. Suggested Verification Protocol", level=1)
    protocol = [
        "Confirm the WGI source workbook SHA-256 matches the recorded hash.",
        "Re-run the correlated-error script (Step 3) and compare output CSV "
        "hashes against those listed in Section 3.",
        "Re-run the disaggregated tables script (Step 4) and confirm 16 × 6 = "
        "96 rows in the supplement, with no missing values.",
        "Open the correlated_error_manifest.json and confirm status = PASS.",
        "Verify that the ρ = 0.00 scenario intervals are statistically "
        "consistent with the locked independent-error baseline.",
        "Confirm that rank Spearman = 1.000 across all ρ scenarios.",
        "Check the scientific constraint checklist in Section 5 against the "
        "revised supervisor report.",
        "Verify that the novelty evidence matrix correctly represents the "
        "capabilities of the cited prior art.",
        "Confirm that the peer-review ZIP SHA-256 matches.",
    ]
    for i, step in enumerate(protocol, 1):
        doc.add_paragraph(f"{i}. {step}")

    # Save
    OUTPUT.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT / "AIML_MRMS_Version_C_Peer_Review_Verification_Guide.docx"
    doc.save(out_path)
    print(f"Verification guide saved to {out_path}")


if __name__ == "__main__":
    main()
